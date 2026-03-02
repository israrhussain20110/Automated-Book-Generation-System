import os
from docx import Document
from core.db_manager import DBManager

class CompilationStage:
    def __init__(self, db: DBManager):
        self.db = db

    def compile_book(self, book_id: str, title: str, output_format: str = "docx"):
        """
        Logic:
        Compile only if:
        final_review_notes_status = no_notes_needed
        OR notes exist for final draft
        Export as .docx or .txt
        """
        chapters = self.db.get_all_chapters(book_id)
        if not chapters:
            print("No chapters found to compile.")
            return None

        if output_format == "docx":
            return self._to_docx(title, chapters)
        else:
            return self._to_txt(title, chapters)

    def _to_docx(self, title, chapters):
        doc = Document()
        doc.add_heading(title, 0)
        
        for ch in chapters:
            doc.add_heading(f"Chapter {ch['chapter_number']}: {ch['title']}", level=1)
            doc.add_paragraph(ch['content'])
            doc.add_page_break()
        
        filename = f"{title.replace(' ', '_')}.docx"
        out_dir = os.path.join(os.path.dirname(__file__), "..", "outputs")
        os.makedirs(out_dir, exist_ok=True)
        out_path = os.path.join(out_dir, filename)
        doc.save(out_path)
        print(f"Book compiled to {out_path}")
        return out_path

    def _to_txt(self, title, chapters):
        filename = f"{title.replace(' ', '_')}.txt"
        out_dir = os.path.join(os.path.dirname(__file__), "..", "outputs")
        os.makedirs(out_dir, exist_ok=True)
        out_path = os.path.join(out_dir, filename)
        with open(out_path, "w") as f:
            f.write(f"{title}\n\n")
            for ch in chapters:
                f.write(f"Chapter {ch['chapter_number']}: {ch['title']}\n")
                f.write(f"{ch['content']}\n\n")
        print(f"Book compiled to {out_path}")
        return out_path
